import docx
import sys

def extract_paragraphs(filepath, label):
    doc = docx.Document(filepath)
    print(f"\n{'='*80}")
    print(f"=== {label}: {filepath} ===")
    print(f"{'='*80}")
    for i, para in enumerate(doc.paragraphs):
        style = para.style.name if para.style else "None"
        text = para.text.strip()
        if not text:
            text = "(empty)"
        # Show first 120 chars
        preview = text[:120] + ("..." if len(text) > 120 else "")
        print(f"\nP{i:03d} | Style: {style}")
        print(f"      | [{preview}]")
    print(f"\nTotal paragraphs: {len(doc.paragraphs)}")

if __name__ == "__main__":
    extract_paragraphs(sys.argv[1], sys.argv[2])

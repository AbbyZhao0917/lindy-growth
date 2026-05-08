#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
淘天案例语言精修脚本
在原始版基础上做语言精修，修改段落文字标蓝，标题不标蓝
"""

import docx
from docx.shared import Pt, RGBColor
from copy import deepcopy

ORIG_PATH = "/home/admin/lindy-growth/cases/taotian_orig.docx"
OUT_PATH = "/home/admin/lindy-growth/cases/taotian_final.docx"

# ============================================================
# REFINED TEXT MAPPING: paragraph_index -> refined_text
# ============================================================

refinements = {}

refinements[7] = (
    '食品安全关乎民生根本，质量提升筑牢产业基石。'
    '当前，食品行业正加速向规模化、连锁化、数字化转型，'
    '生产经营链条持续延伸，监管要素日趋多元，'
    '食品安全管理面临的复杂性与不确定性显著上升。'
    '与此同时，消费者对食品质量、安全及可追溯性的要求不断提高，'
    '倒逼企业持续探索数字技术在食品安全治理中的创新应用。'
)

refinements[8] = (
    '伴随着电商供应链的持续延伸与商品类目的快速更迭，'
    '淘宝天猫平台将人工智能（AI）技术深度融入食品安全与质量管理的核心环节，'
    '构建起一套集\u201c标准引领、智能风控、自动审核、闭环治理\u201d于一体的'
    '多场景数字化食安治理体系。通过前端标准化建设、准入环节AI自动审核、'
    '线上AI问诊及AI客服体验等关键技术，平台实现了从源头准入到末端处置的'
    '全链路智能化升级，切实推动安全管控与品质提升双轮驱动。'
)

refinements[10] = (
    '近年来，人工智能技术在食品安全管理与品质管控领域的应用持续深化，'
    '逐步从辅助工具演变为关键支撑手段。在生产环节，企业广泛运用机器视觉、'
    '智能巡检与预测性维护等技术，对包装缺陷、异物混入、标签错误及设备异常'
    '进行实时识别与预警，显著提升了在线质检与过程控制能力。在流通环节，'
    'AI与物联网、冷链监测、智能仓储等技术深度融合，有效增强了对温湿度变化、'
    '运输异常及库存风险的动态监控能力，进一步提升了食品追溯与风险防控水平。'
    '在餐饮环节，AI技术被广泛应用于后厨环境识别、操作规范监督与门店标准化管理，'
    '推动\u201c明厨亮灶\u201d从可视化展示迈向智能化监管。在电商环节，平台依托AI'
    '开展商品资质审核、违规内容识别、风险商户筛查与异常交易监测，'
    '全面强化了线上食品经营主体的合规治理能力。'
)

refinements[11] = (
    '淘宝天猫平台以撮合交易为基础，连接商家与消费者，构建数字商业生态。'
    '与此同时，平台积极推动产业带合作，链接工厂、产业集群与批发商，'
    '提升供应链效率，满足消费者多元化需求。针对不同业务模式，'
    '平台将AI融入食品安全与品质管理全链路，系统探索电商领域品质管控与效能提升的可行路径。'
)

refinements[13] = (
    '针对产业带中小商家分布分散、标准不统一的痛点，'
    '淘天平台联合百大品牌推出了首个产业带品质标准\u2014\u2014'
    '《淘工厂产业带（国货严选）品质标准》。该标准在严格参照国家标准、'
    '行业标准及平台高标准的基础上融合创新，旨在为新生国货品牌树立品质新标杆。'
)

refinements[14] = (
    '基于该标准，平台建立了涵盖资质审核、源头直采、上架筛查、'
    '实时监控、抽检核查等环节的全链路品控机制，形成系统化的品质保障体系。'
)

refinements[15] = (
    '1. 严格准入与资质审核：运用数字化手段对食品商家生产资质、'
    '商品资质及商品检测报告进行前置审查，从源头把控商家与商品的合规性。'
)

refinements[16] = (
    '2. 源头直采与上架筛查：建立源头追溯机制，商品上架前经多重算法筛查，'
    '拦截潜在风险品，并依据风险等级开展实地审查与验证。'
)

refinements[17] = (
    '3. 实时监控与熔断机制：对在售商品实施全天候数据监控，'
    '结合随机双盲抽检机制，当检测指标出现异常或存在风险时，'
    '系统即刻触发下架熔断流程。'
)

refinements[18] = (
    '4. 售后响应与逆向追溯：建立快速响应通道，确保问题食品可追溯、可召回。'
    '针对产业带中小商家质量管理与运营能力偏弱的问题，平台推出智能AI客服系统。'
    '该系统运用自然语言处理技术，自动分析高频投诉点，反向推导产品潜在的品质缺陷'
    '或描述不符问题，形成\u201c消费端反馈\u2192生产端改进\u201d的质量闭环，'
    '有效助力中小商家提升品控意识与服务水平。'
)

refinements[20] = (
    '在平台自营模式下，天猫超市针对食品安全的高敏感性，'
    '研发了\u201c品质管控AI智能体\u201d，将原本依赖人工经验的非标准化工作'
    '全面升级为AI驱动的标准化作业，重点聚焦于风险前置识别与供应商主动合规管理。'
)

refinements[22] = (
    '构建\u201c感知\u2014分析\u2014诊断\u2014处置\u201d智能引擎，'
    '搭建三大核心功能模块，实现品质治理的智能化闭环。一是AI问数，'
    '改变传统人工跨看板调取食品安全数据的低效模式，通过自然语言描述即可在秒级'
    '完成数据提取、清洗与聚合，实现食品安全管理数据的即时洞察。二是品质问诊，'
    '基于历史食品安全案例库与规则知识图谱，AI自动关联用户评价、售后记录等多维信息，'
    '输出结构化归因报告，并智能推荐下架、整改或处罚等处置方案。三是品质巡检，'
    '根据商品风险等级配置自动化巡检规则，周期性推送品质报告，'
    '实现对食品安全风险的常态化、主动式监控。'
)

refinements[24] = (
    '打造高效可信的智能审核防线。针对商品标签、仓储环境及商家资质等高频率审核场景，'
    '部署AI自动审核系统。在标签与资质审核方面，利用图像识别技术自动核查食品包装'
    '标签合规性，包括配料表、营养成分表及商家经营资质有效性，确保各项指标符合要求。'
    '在仓储环境审核方面，通过智能视觉技术分析仓储监控画面，实时识别温度异常及环境、'
    '操作等违规风险，保障食品存储环节的安全。该系统不仅提升了审核效率，还统一了处置标准，'
    '保障了合规治理的公平性。'
)

refinements[26] = (
    '为推动从\u201c被动整改\u201d向\u201c主动合规\u201d转变，'
    '解决供应商因不熟悉规则而反复违规的问题，平台建立了AI驱动的供应商成长体系。'
    '具体而言，AI主动识别风险数据并推送规则培训；基于商家销售类目及历史违规数据，'
    '精准识别合规薄弱点，定向推送轻量化、场景化的学习内容，有效降低首次违规率。'
    '此外，当食品安全违规行为发生时，AI提供针对性的整改指导方案，'
    '帮助商家快速定位问题根源，缩短整改周期。'
)

refinements[28] = (
    '借助AI技术赋能，平台商家在食品安全管控与品质提升方面取得了显著成效。'
    '首先，风险防控前置，筑牢安全防线。通过AI全链路管控技术，平台改变了传统依赖'
    '事后处置的被动局面，实现了从\u201c事后补救\u201d到\u201c事前预防\u201d的转变。'
    '基于AI主动感知与早期预警，平台成功拦截大量潜在食品安全隐患，'
    '有效防范了因安全问题引发的舆情发酵及风险事件。同时，依托数据的动态监控机制，'
    '平台显著降低了食品安全事故的实际发生率，维护了平台声誉与市场秩序。'
)

refinements[29] = (
    '其次，治理效能全面升级，实现降本增效。依托标准化体系搭建、AI问诊、'
    '自动审核及实时巡检等技术手段，平台实现了食品安全风险的'
    '\u201c早发现、早预警、早处置\u201d。数据获取效率提升70%以上，'
    '问题发现时效平均缩短超过50%，问题诊断准确率高达98%，'
    '有效避免了因人为疏忽导致的食品安全漏检。'
    '商家首次违规率降低27%，重复违规率降低67%，系统性食品安全风险显著下降。'
)

refinements[30] = (
    '第三，有效扶持商家，构建共建共治生态。平台通过优化规则引导、'
    '强化能力赋能与提供精准服务，帮助商家在提升食品安全管理水平的同时降低合规成本。'
    '依托AI问诊、智能审核、风险预警等工具，平台为商家提供经营过程中的实时提示与整改建议，'
    '使其能够更便捷地识别问题、纠正偏差，提升自主治理能力。对于新入驻商家与中小商家，'
    '平台通过标准化培训、操作指引与案例宣导，强化其食品安全意识与规范经营能力，'
    '降低因经验不足带来的违规风险。通过\u201c技术支持+规则引导+能力提升\u201d的综合策略，'
    '平台不仅提高了商家的合规水平，也促进了平台与商家之间的协同联动，'
    '逐步形成了共建、共治、共享的食品安全治理生态。'
)


def is_title_paragraph(para):
    """Determine if a paragraph is a title that should not be colored blue."""
    text = para.text.strip()
    if not text:
        return True
    
    # Title patterns
    if text.startswith(('\u4e00\u3001', '\u4e8c\u3001', '\u4e09\u3001',
                        '\uff08\u4e00\uff09', '\uff08\u4e8c\uff09',
                        '\u9644\u4ef6', '\u667a\u9a6d', '2025', '*')):
        return True
    
    # Sub-headings
    if text in ('1.AI\u95ee\u8bca', '2.AI\u5ba1\u6838', '3.AI\u9a71\u52a8'):
        return True
    
    # Check if all runs are bold (title style)
    if para.runs and all(r.bold for r in para.runs if r.text.strip()):
        return True
    
    return False


def copy_and_refine():
    doc = docx.Document(ORIG_PATH)
    
    for para_idx, para in enumerate(doc.paragraphs):
        if para_idx not in refinements:
            continue
        
        refined_text = refinements[para_idx]
        original_text = para.text
        
        if not original_text.strip():
            continue
        
        is_title = is_title_paragraph(para)
        
        # Get original font properties from first run
        first_run = para.runs[0] if para.runs else None
        orig_size = first_run.font.size if first_run else None
        orig_name = first_run.font.name if first_run else None
        
        # Remove all existing runs
        runs_to_remove = list(para.runs)
        for r in runs_to_remove:
            r._element.getparent().remove(r._element)
        
        # Add new run with refined text
        new_run = para.add_run(refined_text)
        
        if orig_size:
            new_run.font.size = orig_size
        if orig_name:
            new_run.font.name = orig_name
        
        if is_title:
            # Titles: preserve bold and color, no blue
            new_run.bold = True
            # Keep original dark color (333333 for section headers)
            new_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        else:
            # Content paragraphs: color blue
            new_run.bold = None
            new_run.font.color.rgb = RGBColor(0, 102, 204)
    
    doc.save(OUT_PATH)
    print(f"Saved to {OUT_PATH}")
    print("Done!")


if __name__ == "__main__":
    copy_and_refine()

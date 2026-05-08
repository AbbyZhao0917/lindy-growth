#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
瑞幸咖啡语言精修脚本
基于ruixing_orig.docx，对段落文字做语言精修并标蓝，输出ruixing_final.docx
"""
from docx import Document
from docx.shared import RGBColor
import os

BLUE = RGBColor(0x2E, 0x75, 0xB6)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_PATH = os.path.join(BASE_DIR, 'ruixing_orig.docx')
OUTPUT_PATH = os.path.join(BASE_DIR, 'ruixing_final.docx')

doc = Document(INPUT_PATH)

# Define refinements per paragraph index
# Rules: 去口语化、冗余精简、动词精准化、数据表述规范化、增加过渡引导
# Headings (P1, P4, P5, P10, P16, P21, P23, P29, P34) are NOT modified

refinements = {
    # P2 - First paragraph of case background
    2: (
        '在健康消费理念持续深化的背景下，消费者对现制饮品的诉求已从"口感偏好"'
        '转向"成分安全、配方透明、原料优质可溯"。然而，食品安全与健康领域的'
        '热点问题频发——人工合成甜味剂阿斯巴甜被国际癌症研究机构（IARC）列为'
        '2B类"可能致癌物"，部分氢化油被世界卫生组织列为心血管健康"一级风险'
        '因子"——这些事件加剧了消费者的隐忧。'
    ),
    
    # P3 - Second paragraph of case background
    3: (
        '与此同时，国家"十四五"国民健康规划等政策相继出台，最新标准法规持续'
        '抬升行业准入门槛。作为连锁饮品行业首家突破万店规模的头部品牌，瑞幸咖啡'
        '主动扛起品质引领责任，以"品质至上"为核心，在全价值链食品安全管理中聚焦'
        '原料风险物质拦截。通过构建"合规标准+内控标杆"双维管理体系，将消费者的'
        '"健康关切"转化为"企业内控标准"，以标准化手段从源头阻断风险物质进入消费端，'
        '并依托数字化工具实现管控效率升级，真正为消费者筑牢"成分安心、饮用放心"的'
        '健康防线。'
    ),
    
    # P6 - First paragraph under "率先建立禁用物质清单"
    6: (
        '为树立行业品质标杆，瑞幸咖啡于2024年率先发布《瑞幸咖啡配料使用指引》，'
        '构筑以"禁用物质清单"为核心的企业级配料风险拦截体系。该体系覆盖逾百种潜在'
        '高风险物质，管控范围远超国家强制标准，是行业内首个全面实施的"严格于基准限量"'
        '线上配料管控系统，以"企业内控金线"重构饮品配料安全新基准。'
    ),
    
    # P7 - Second paragraph under the same section
    7: (
        '该指引以"风险评估—分级管控—线上监控—动态进化"为逻辑主线，构建覆盖'
        '"公众认知风险""健康危害风险""监管合规风险""清洁配方需求"的多维筛选体系。'
        '综合考量国内外风险物质评估数据，划定"禁止使用+禁限使用"双级管控清单，'
        '其中核心禁用项涵盖植脂末、部分氢化油等70余种风险物质及10余种潜在争议成分。'
    ),
    
    # P8 - Third paragraph - 清洁配方工程
    8: (
        '在具体执行中，以"清洁配方工程"为例，瑞幸通过"禁用物质清单+定制原料升级"'
        '进行原料管理，叠加线上"配方审核"进行配方管理，三线并行推进产品清洁配方升级。'
        '具体而言，将风险添加剂列入"禁用物质清单"，从源头上阻断其使用；在研发阶段参与'
        '原料定制，遵循"配料简单纯粹"原则优化核心原料；同时建立前置审核机制，配方须经'
        '线上系统预审，确认符合国家标准添加剂规定及"禁用物质清单"要求，系统自动对风险'
        '物质发出预警。这一将清洁配方理念融入全流程管控的实践，使产品在研发阶段即具备'
        '"成分透明、天然健康"的属性，有效降低消费者对配料安全的疑虑。'
    ),
    
    # P11 - 数字化运用 paragraph
    11: (
        '瑞幸咖啡打造了数字化审核体系，依托数字化系统实现全价值链风险物质智能管控与'
        '前置预警。'
    ),
    
    # P13 - 动态标法中枢 paragraph
    13: (
        '该系统深度集成GB 2760等国内外法规数据库及《瑞幸咖啡配料使用指引》等内控要求，'
        '自动解析食品添加剂限量标准与风险物质指标，形成动态数据库。当法规更新时（如新增'
        '某食品类别防腐剂限量），数字化系统实时同步完成全量配方扫描，确保新老产品配方'
        '全面合规。'
    ),
    
    # P15 - 数据穿透式核验 paragraph
    15: (
        '通过系统对接，自动抓取原料及饮品配方信息，基于预设的系统规则自动校验饮品添加剂'
        '及原料配方的合规性。'
    ),
    
    # P18 - 超标即时拦截 paragraph
    18: (
        '当原料配方中出现风险物质，或饮品添加剂使用量超过国家标准限量时，系统立即阻断'
        '相关流程，并自动生成风险溯源报告进行预警。'
    ),
    
    # P20 - 阈值预警监测 paragraph
    20: (
        '当使用量触及国家标准阈值时，系统自动将该饮品纳入"食品安全风险监测体系"，进入'
        '下一环节的动态监控与外检验证。'
    ),
    
    # P22 - Opening paragraph of 成果和价值 section
    22: (
        '瑞幸咖啡通过建立《瑞幸咖啡配料使用指引》并融合数字化应用，实现了风险物质的'
        '自动化智能管控与前置预警，在行业标准重构、管理效能跃升、消费信任夯实及质量'
        '管理数字化转型方面取得了突破性成果：'
    ),
    
    # P25 - 超国标管控 paragraph
    25: (
        '依托覆盖逾100种风险物质的管控清单，标准严格度达到行业领先水平。《瑞幸咖啡配料'
        '使用指引》实施以来，已成功拦截多次管控物质违规，从源头上杜绝了合规风险及舆情争议。'
    ),
    
    # P27 - 全价值链管控 paragraph
    27: (
        '在研发阶段前置产品配方审核管控，确保新配方符合"清洁配方"升级要求，同步推动'
        '上游供应链完成原料清洁化升级，形成"需求—供给"的品质闭环。以瑞幸咖啡"果C美式"'
        '系列为例，通过工艺与配方双升级优化主要原料配方，实现了清洁配方从标准框架到产品'
        '落地的完整闭环。'
    ),
    
    # P31 - 实现「算法替代人工」的效率革命 paragraph
    31: (
        '结合数字化审核体系，风险物质合规审核响应时间从天级压缩至小时级，审核人力成本'
        '节省约40%，审核时效提升约80%。'
    ),
    
    # P33 - 双级风控推动行业监管模式升级 paragraph
    33: (
        '通过设置"超标即时拦截+阈值预警监测"机制，实现了风险物质管理从"事后抽检"到'
        '"实时监控"的跨越：当监控物质用量触及阈值时，系统自动联动门店端外检，推动现制'
        '饮品向"智能监管"模式转型。'
    ),
    
    # P36 - 消费信任具象化 paragraph
    36: (
        '基于配方及原料信息的系统化与可追溯性，瑞幸咖啡推进"成分透明+数据量化"策略，'
        '将相关营养信息标注于饮品包装杯身，实现产品身份信息透明化。消费者可清晰了解饮品'
        '的营养信息，精准匹配自身控糖、控卡等健康化需求，真正实现"喝得明白、选得安心"。'
    ),
    
    # P38 - 行业生态引领 paragraph
    38: (
        '在建立清洁配料指引的基础上，瑞幸咖啡进一步推出"三个100"轻乳茶标准，推动上下游'
        '建立"零植脂末、零反式脂肪酸"等六项内控标准，促进上游供应商工艺升级，形成从原料'
        '到终端的品质保障体系，为新消费品牌树立了可复制的标杆范式。'
    ),
    
    # P39 - 总结 paragraph
    39: (
        '综上所述，瑞幸咖啡通过构建配料使用指引与数字化审核体系的有机融合，在食品安全治理'
        '与科技创新层面取得了显著成效，为连锁饮品行业的健康升级提供了一条兼具战略前瞻性与'
        '实操性的示范路径。'
    )
}

def set_paragraph_text_blue(para, new_text):
    """Replace paragraph text with blue-colored text, preserving paragraph style"""
    # Clear existing runs - set their text to empty
    for run in para.runs:
        run.text = ''
    
    # Set text in first run or add new run
    if para.runs:
        run = para.runs[0]
        run.text = new_text
        run.font.color.rgb = BLUE
    else:
        run = para.add_run(new_text)
        run.font.color.rgb = BLUE

# Apply refinements
updated = []
for p_idx, new_text in sorted(refinements.items()):
    if p_idx < len(doc.paragraphs):
        para = doc.paragraphs[p_idx]
        old_text = para.text[:60] + '...' if len(para.text) > 60 else para.text
        set_paragraph_text_blue(para, new_text)
        updated.append(f"P{p_idx}: {old_text}")
        print(f"P{p_idx} updated ✓")
    else:
        print(f"P{p_idx} NOT FOUND ✗")

doc.save(OUTPUT_PATH)
print(f"\nSaved: {OUTPUT_PATH}")
print(f"Total paragraphs refined: {len(updated)}")

# Verify
doc2 = Document(OUTPUT_PATH)
print("\n=== Verification ===")
for i, p in enumerate(doc2.paragraphs):
    has_blue = False
    for r in p.runs:
        if r.font.color and r.font.color.rgb == BLUE:
            has_blue = True
            break
    if has_blue:
        print(f"P{i} [BLUE]: {p.text[:80]}...")
    elif p.text.strip():
        style = p.style.name
        is_heading = 'Heading' in style or '标题' in style or 'head' in style.lower()
        print(f"P{i} [{'HEAD' if is_heading else 'BLACK'}]: {p.text[:80]}...")

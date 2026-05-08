#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
运荔枝案例文档语言精修脚本
基于原始版 yunlizhi_orig.docx 进行精修，修改段落文字标蓝
输出 yunlizhi_final.docx

精修方向：
1. 去口语化
2. 冗余精简
3. 动词精准化
4. 数据表述规范化
5. 增加过渡引导
"""

from docx import Document
from docx.shared import RGBColor

BLUE = RGBColor(0, 0, 255)

def set_para_text(para, text, bold=None):
    for run in para.runs:
        run.text = ''
    if para.runs:
        para.runs[0].text = text
        if bold is not None:
            para.runs[0].bold = bold
    else:
        r = para.add_run(text)
        if bold is not None:
            r.bold = r

def make_blue(para):
    for run in para.runs:
        run.font.color.rgb = BLUE

def set_and_blue(para, text, bold=None):
    set_para_text(para, text, bold)
    make_blue(para)

LQ = '\u201c'  # Left curly quote "
RQ = '\u201d'  # Right curly quote "

doc = Document('/home/admin/lindy-growth/cases/yunlizhi_orig.docx')
paras = doc.paragraphs
print(f"Total paragraphs: {len(paras)}")

edits = {}

# P0-P13: Form header and template content — delete
for idx in range(0, 14):
    edits[idx] = ""

# P14: Title enhancement
edits[14] = f"运荔枝：打造低温优质乳冷链物流技术规范"

# P15: "一、案例背景" — title, keep unchanged
# P16: Remove instruction text
edits[16] = ""

# P28: Remove standalone quote attribution
edits[28] = ""

# P29: Background intro — remove "近年来", tighten
edits[29] = (
    f"随着消费者对乳制品{LQ}新鲜、营养、安全{RQ}需求的持续提升，"
    f"低温鲜奶（巴氏杀菌乳）市场迎来爆发式增长。"
    f"然而，低温鲜奶作为典型的高温敏感商品，其品质高度依赖从{LQ}牧场到餐桌{RQ}全链条的冷链完整性。"
    f"当前行业普遍面临以下挑战："
)

# P35: Industry background — refine
edits[35] = (
    f"2025年正值国家优质乳工程实施十周年。目前，优质乳工程已覆盖全国29个省（自治区、直辖市）的"
    f"79家乳企，2024年优质巴氏杀菌乳产量占全国巴氏杀菌乳总量比例达97%。"
    f"以新希望乳业为代表的优秀品牌推动优质鲜乳全面崛起，国产奶业进入{LQ}鲜活时代{RQ}。"
)

# P36: Remove instruction text
edits[36] = ""

# P39: Remove standalone action title
edits[39] = ""

# P42: Company response — add transition, refine
edits[42] = (
    f"针对上述行业痛点，鲜生活冷链旗下数智物流平台运荔枝，依托十年优质乳冷链服务经验及数智化技术沉淀，"
    f"国内首发覆盖{LQ}牧场—工厂—仓配—门店{RQ}全场景的《优质乳冷链物流技术规范》企业标准"
    f"（标准编号：Q/CXSH010—2025），旨在通过标准化、数智化手段系统解决低温优质乳冷链物流中的痛点，"
    f"推动行业从{LQ}粗放竞争{RQ}转向{LQ}精益运营{RQ}。"
)

# P44: Remove instruction text
edits[44] = ""

# P46: Quote about milk — refine
edits[46] = (
    f"好牛奶的最大营养价值在于其{LQ}天然鲜活营养{RQ}，"
    f"无论是生乳还是成品奶，品质保障均离不开专业运输车辆。"
)

# P47: Vehicle准入 — minor refine
edits[47] = (
    f"运荔枝旗下云牧安达作为专业农牧物流服务商，对原奶运输车辆实施严格准入管理："
)

# P49-P54: Keep as technical specs
# P58: City distribution — refine
edits[58] = (
    f"城配运输由云海店配负责落地，采用四星级以上车辆标准，100%设备智能应用，360°无死角监控。"
    f"同时，积极推行低碳冷链解决方案，推广新能源冷藏车应用，有效实现绿色减排，"
    f"以实际行动践行ESG可持续发展理念。"
)

# P61: Operation standards intro — refine
edits[61] = (
    f"针对运输环节中常见的{LQ}断链{RQ}风险，运荔枝构建了9项关键时效节点与30项操作规范，"
    f"覆盖{LQ}牧场—工厂—仓配—门店{RQ}全场景，明确温度标准、时效标准及预警机制，"
    f"确保各环节有章可循，有效降低乳品损耗。"
)

# P74: Auto control description — refine
edits[74] = (
    f"所有节点均在运荔枝数智化系统中实现自动卡控与预警，"
    f"任一环节触发预警即启动分级响应，确保实时干预，杜绝断链风险。"
)

# P77: Product matrix intro — minor refine
edits[77] = (
    f"运荔枝依托数智化底座，打造了覆盖{LQ}指引—监管—预警—卡控{RQ}全流程的数智产品矩阵，核心产品包括："
)

# P78: 荔途守望 — refine verbose expressions
edits[78] = (
    f"荔途守望（物流大屏）\n"
    f"数智产品{LQ}荔途守望{RQ}提供物流大屏、运途管理、异常监控、数据分析四大核心功能。"
    f"客户可通过该产品实时查看服务全流程质量、监管服务过程、追溯异常信息，"
    f"实现服务透明化管理，高效规避运途中食安风险，确保每车货物准时必达。"
)

# P90: Remove raw title
edits[90] = ""

# P93: Results data — add transition, refine
edits[93] = (
    f"经过系统建设，运荔枝的标准化与数智化体系取得了扎实的成效。"
    f"针对低温优质乳的标准化建设不仅推动自身提质增效，更赋能行业从{LQ}粗放竞争{RQ}转向{LQ}精益运营{RQ}。"
    f"自标准实施以来，到仓及时率稳定在99.5%以上，到店及时率最高达99.9%，"
    f"终端配送时效管控稳定可靠；温度达标率在数智化系统辅助下稳定在99.9%以上。"
)

# P95-P98: Innovation items — merge
edits[95] = (
    f"在创新层面，该标准体现出四个显著特征："
    f"一是标准首创性与系统性，国内首个覆盖{LQ}牧场—工厂—仓配—门店{RQ}全场景的优质乳冷链物流企业标准，填补了行业空白；"
    f"二是标准与数智化深度绑定，将30项操作规范、9大时效节点、9项数智功能固化到系统流程，"
    f"实现{LQ}标准即流程、流程即卡控{RQ}，杜绝人为偏差；"
    f"三是分级预警与闭环管理，建立四级预警机制（温度、时效、设备、合规），"
    f"配套应急预案与处置SOP，实现{LQ}异常早发现、快处置、可追溯{RQ}；"
    f"四是绿色低碳融合，推广新能源冷藏车，量化碳排放数据，将ESG理念融入冷链标准。"
)
edits[96] = ""
edits[97] = ""
edits[98] = ""

# P100-P102: Promotion value items — merge
edits[100] = (
    f"该标准已在行业内展现出良好的推广复制价值。"
    f"目前已在新希望乳业等头部乳企全面落地，并向中小乳企开放技术规范与数智化工具，"
    f"助力行业整体提升冷链水平。"
)
edits[101] = (
    f"在跨品类应用方面，标准中的温控、时效、追溯等核心模块可平移至冰品、鲜肉、鲜食等短保食品品类，"
    f"具备较强的横向复制能力。"
)
edits[102] = (
    f"此外，该标准与国家优质乳工程、食品安全{LQ}两个责任{RQ}等政策高度契合，"
    f"可作为行业团体标准或国家标准的前期实践基础。"
)

# P103: Closing summary — refine
edits[103] = (
    f"运荔枝通过{LQ}标准+数智+运营{RQ}三位一体的模式，"
    f"不仅实现了自身服务能力的跃升，更形成了可输出、可复制的冷链物流创新解决方案，"
    f"为食品行业高质量发展提供了坚实底座。"
)

# Apply all edits
# Headings that should NOT be blued:
# P15 一、案例背景
# P43 二、实施过程与路径
# P45 (一）运输硬件标准
# P59 (二）30项操作规范与9大时效节点
# P62 生乳运输关键操作规范（部分）
# P67 成品奶运输关键操作规范（部分）
# P72 9大时效节点与预警机制
# P76 (三）9项数智产品
# P85 任务推送与卡控
# P91 (一）量化成效数据
# P94 (二）创新点提炼
# P99 (三）推广复制价值
# P90 成果和价值 — will be emptied, so no issue

heading_indices = {15, 43, 45, 59, 62, 67, 72, 76, 85, 91, 94, 99}

for idx, text in edits.items():
    if idx < len(paras):
        p = paras[idx]
        if idx in heading_indices:
            set_para_text(p, text)
            if text:
                print(f"Set heading P{idx}: {text[:60]}...")
            else:
                print(f"Emptied heading P{idx}")
        else:
            set_and_blue(p, text)
            if text:
                print(f"Edited P{idx}: {text[:60]}...")
            else:
                print(f"Emptied P{idx}")

output_path = '/home/admin/lindy-growth/cases/yunlizhi_final.docx'
doc.save(output_path)
print(f"\nSaved to {output_path}")
print("Done!")

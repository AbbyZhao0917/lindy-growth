#!/usr/bin/env python3
"""
海天味业 docx 语言精修脚本
修改项：
1. 全文统一中英文标点：英文逗号→中文顿号或中文逗号
2. 口语化开头修改：「另外海天深耕...」→「海天深耕...」
3. 冗余表述精简：「明显大幅提升」→「大幅提升」
4. 数据补充：模糊数据「成本节约1000万以上」→「成本节约超1000万元」
5. 弱表述强化：「在一定程度上成功解决了」→「有效解决了」
6. 句式补全：引号开头的段落补主语
7. 主观评价克制：「陆续获得」→「获评」
"""

from docx import Document
from docx.shared import RGBColor

BLUE = RGBColor(0, 0, 255)

def main():
    input_path = "haitian_orig.docx"
    output_path = "haitian_final.docx"
    
    doc = Document(input_path)
    
    # ==========================================
    # Para 12: Fix English commas → Chinese commas/顿号
    # ==========================================
    para = doc.paragraphs[12]
    for run in para.runs:
        old = run.text
        new = old
        # Specific replacements for para 12
        new = new.replace('投入,建成SAP,MES,APS,WMS等系统', '投入，建成SAP、MES、APS、WMS等系统')
        new = new.replace('边界,实现生产过程', '边界，实现生产过程')
        new = new.replace('云计算，大数据,5G等', '云计算、大数据、5G等')
        new = new.replace('工具,打造数据', '工具，打造数据')
        # Also handle the case where parts of these strings might be split across runs
        new = new.replace(',建成SAP,MES,APS,WMS', '，建成SAP、MES、APS、WMS')
        new = new.replace(',MES,APS,WMS等系统', '、MES、APS、WMS等系统')
        new = new.replace(',APS,WMS等系统', '、APS、WMS等系统')
        new = new.replace(',WMS等系统', '、WMS等系统')
        if new != old:
            run.text = new
            run.font.color.rgb = BLUE
    
    # ==========================================
    # Para 13: 「另外海天深耕...」→「海天深耕...」
    # ==========================================
    para = doc.paragraphs[13]
    for run in para.runs:
        old = run.text
        new = old.replace('另外海天深耕', '海天深耕')
        if new != old:
            run.text = new
            run.font.color.rgb = BLUE
    
    # ==========================================
    # Para 48: 「明显大幅提升」→「大幅提升」 + 「成本节约1000万以上」→「成本节约超1000万元」
    # ==========================================
    para = doc.paragraphs[48]
    for run in para.runs:
        old = run.text
        new = old.replace('明显大幅提升', '大幅提升')
        new = new.replace('成本节约1000万以上', '成本节约超1000万元')
        if new != old:
            run.text = new
            run.font.color.rgb = BLUE
    
    # ==========================================
    # Para 49: 「陆续获得」→「获评」
    # ==========================================
    para = doc.paragraphs[49]
    for run in para.runs:
        old = run.text
        new = old.replace('陆续获得', '获评')
        if new != old:
            run.text = new
            run.font.color.rgb = BLUE
    
    # ==========================================
    # Para 53: 「在一定程度上成功解决了」→「有效解决了」 + 引号补主语
    # 「"品控智能化技术"是...」→「海天"品控智能化技术"是...」
    # ==========================================
    para = doc.paragraphs[53]
    # Check run 0: contains "品控智能化技术"是
    # We need to add 海天 at the beginning
    if para.runs[0].text.startswith('"品控智能化技术"是') or \
       para.runs[0].text.startswith('\u201c品控智能化技术\u201d是'):
        old = para.runs[0].text
        para.runs[0].text = '海天' + old
        para.runs[0].font.color.rgb = BLUE
    
    for run in para.runs:
        old = run.text
        new = old.replace('在一定程度上成功解决了', '有效解决了')
        if new != old:
            run.text = new
            run.font.color.rgb = BLUE
    
    doc.save(output_path)
    print(f"Saved to {output_path}")
    
    # Verify all changes
    doc2 = Document(output_path)
    checks = [
        (12, '投入，建成SAP、MES、APS、WMS', '英文逗号→中文逗号/顿号'),
        (13, '海天深耕酱油科技领域多年', '去掉「另外」'),
        (48, '大幅提升', '冗余精简'),
        (48, '成本节约超1000万元', '数据补充'),
        (49, '获评', '主观评价克制'),
        (53, '海天"品控智能化技术"', '句式补全(主语)'),
        (53, '有效解决了', '弱表述强化'),
    ]
    all_ok = True
    for i, needle, desc in checks:
        text = doc2.paragraphs[i].text
        if needle in text:
            print(f"  ✓ [{i}] {desc}: found '{needle}'")
        else:
            print(f"  ✗ [{i}] {desc}: NOT found - '{needle}'")
            print(f"    Actual: {text[:150]}...")
            all_ok = False
    
    if all_ok:
        print("\nAll modifications applied successfully!")
    
    # Check blue coloring
    print("\n=== Blue-colored runs verification ===")
    for i in [12, 13, 48, 49, 53]:
        para = doc2.paragraphs[i]
        blue_count = 0
        for run in para.runs:
            if run.font.color and run.font.color.rgb and str(run.font.color.rgb) == '0000FF':
                blue_count += 1
        print(f"  Para {i}: {blue_count} blue runs (text: {para.text[:80]}...)")

if __name__ == '__main__':
    main()
